"""
DOCX Reconstruction Module
===========================
Builds a Word document from OCR + translation results using python-docx.

Layout approach
---------------
DOCX does not support absolute pixel positioning portably, so this builder
uses a flowing, readable layout:

  • One section per page, with a page-break between pages.
  • A reduced-size thumbnail of the original page is inserted at the top.
  • For each OCR text block the builder emits either:
      – A plain paragraph (text was already in English / not Russian), or
      – A side-by-side annotation: the Russian original in grey italic on the
        first line, the English translation in normal black on the next, with
        a light-blue left border to make translated blocks visually distinct.
  • A JSON metadata appendix is added at the end.
"""

from __future__ import annotations

import io
import json
import logging
from typing import Any, Dict, List, Optional

from PIL import Image as PILImage

logger = logging.getLogger(__name__)

# Thumbnail max width in inches inserted above each page's text
_THUMB_WIDTH_INCHES = 5.0

# Indentation for translated-block annotation (in inches)
_INDENT_INCHES = 0.2


def _inches(val: float) -> Any:
    from docx.shared import Inches
    return Inches(val)


def _pt(val: float) -> Any:
    from docx.shared import Pt
    return Pt(val)


def _rgb(r: int, g: int, b: int) -> Any:
    from docx.shared import RGBColor
    return RGBColor(r, g, b)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _insert_thumbnail(doc: Any, page_img: PILImage.Image) -> None:
    """Insert a resized thumbnail of the page, centred."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        buf = io.BytesIO()
        thumb = page_img.copy()
        # Scale to at most _THUMB_WIDTH_INCHES wide
        ratio = (_THUMB_WIDTH_INCHES * 96) / thumb.width   # 96 px/in screen
        new_w = int(thumb.width  * ratio)
        new_h = int(thumb.height * ratio)
        thumb = thumb.resize((new_w, new_h), PILImage.LANCZOS)
        thumb.save(buf, format="PNG")
        buf.seek(0)

        p = doc.add_picture(buf, width=_inches(_THUMB_WIDTH_INCHES))
        # Centre the paragraph that contains the picture
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore[attr-defined]
    except Exception as exc:
        logger.warning("Could not insert thumbnail: %s", exc)


def _add_translated_block(
    doc: Any,
    original: str,
    translated: str,
    confidence: float,
) -> None:
    """
    Emit one translated text block as two stacked paragraphs:
      Line 1 – Russian original  (grey italic, small)
      Line 2 – English translation (normal black, slightly larger)
    with a subtle left border to signal "this was translated".
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Russian original ────────────────────────────────────────────────────
    orig_para = doc.add_paragraph()
    orig_para.paragraph_format.left_indent  = _inches(_INDENT_INCHES)
    orig_para.paragraph_format.space_after  = _pt(0)
    orig_para.paragraph_format.space_before = _pt(2)

    run_orig = orig_para.add_run(f"[RU] {original}")
    run_orig.italic          = True
    run_orig.font.size       = _pt(9)
    run_orig.font.color.rgb  = _rgb(0x88, 0x88, 0x88)

    # ── English translation ─────────────────────────────────────────────────
    trans_para = doc.add_paragraph()
    trans_para.paragraph_format.left_indent  = _inches(_INDENT_INCHES)
    trans_para.paragraph_format.space_after  = _pt(4)
    trans_para.paragraph_format.space_before = _pt(0)

    run_trans = trans_para.add_run(translated)
    run_trans.font.size      = _pt(11)
    run_trans.font.color.rgb = _rgb(0x0D, 0x0D, 0x0D)

    # ── Add left border to both paragraphs (blue bar visual cue) ───────────
    for para in (orig_para, trans_para):
        try:
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"),   "single")
            left.set(qn("w:sz"),    "4")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "4472C4")   # Word "Blue Accent 1"
            pBdr.append(left)
            pPr.append(pBdr)
        except Exception:
            pass  # border is cosmetic — skip silently


def _add_plain_block(doc: Any, text: str) -> None:
    """Emit a plain (untranslated) text block as a normal paragraph."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after  = _pt(4)
    para.paragraph_format.space_before = _pt(2)
    for run in para.runs:
        run.font.size = _pt(11)


def _add_metadata_appendix(doc: Any, pages_data: List[Dict[str, Any]]) -> None:
    """Append a JSON metadata section at the end of the document."""
    doc.add_page_break()
    doc.add_heading("OCR Metadata (JSON)", level=2)

    payload = {
        "pages": [
            {
                "page_number": p.get("page_number"),
                "word_count": p.get("word_count", 0),
                "translation_word_count": p.get("translation_word_count", 0),
                "num_columns": p.get("num_columns", 1),
                "blocks": [
                    {
                        "bbox": b.get("bbox"),
                        "original_text": b.get("text"),
                        "translated_text": b.get("translated_text"),
                        "confidence": round(b.get("confidence", 0.0), 2),
                        "was_translated": b.get("was_translated", False),
                    }
                    for b in p.get("blocks", [])
                ],
            }
            for p in pages_data
        ]
    }
    json_str = json.dumps(payload, ensure_ascii=False, indent=2)

    para = doc.add_paragraph(json_str)
    for run in para.runs:
        run.font.name = "Courier New"
        run.font.size = _pt(7)
        run.font.color.rgb = _rgb(0x44, 0x44, 0x44)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_docx(
    pages_data: List[Dict[str, Any]],
    page_images: List[PILImage.Image],
    include_thumbnails: bool = True,
) -> bytes:
    """
    Build a .docx document from OCR / translation results.

    Parameters
    ----------
    pages_data        : Per-page dicts from process_page().
    page_images       : Original PIL images, one per page.
    include_thumbnails: Insert page thumbnails above each page's text.

    Returns
    -------
    Raw bytes of the .docx file.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise ImportError("python-docx is required: pip install python-docx") from exc

    doc = Document()

    # ── Document title ───────────────────────────────────────────────────────
    title = doc.add_heading("OCR Translation Result", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore[attr-defined]

    for page_idx, (page_data, page_img) in enumerate(zip(pages_data, page_images)):
        page_num   = page_data.get("page_number", page_idx + 1)
        blocks     = page_data.get("blocks", [])
        num_cols   = page_data.get("num_columns", 1)
        proc_time  = page_data.get("processing_time", 0.0)
        word_count = page_data.get("word_count", 0)
        trans_count= page_data.get("translation_word_count", 0)

        if page_idx > 0:
            doc.add_page_break()

        # ── Page heading ─────────────────────────────────────────────────────
        doc.add_heading(f"Page {page_num}", level=2)

        # ── Metadata line ────────────────────────────────────────────────────
        meta = doc.add_paragraph(
            f"Words: {word_count}  •  Translated: {trans_count}  •  "
            f"Columns: {num_cols}  •  Time: {proc_time:.2f} s"
        )
        for run in meta.runs:
            run.font.size = _pt(9)
            run.italic    = True
            run.font.color.rgb = _rgb(0x88, 0x88, 0x88)
        meta.paragraph_format.space_after = _pt(6)

        # ── Page thumbnail ────────────────────────────────────────────────────
        if include_thumbnails and page_img is not None:
            _insert_thumbnail(doc, page_img)
            doc.add_paragraph()   # spacer

        # ── Text blocks ───────────────────────────────────────────────────────
        doc.add_heading("Extracted Text", level=3)

        if not blocks:
            doc.add_paragraph("(No text detected on this page.)")
        else:
            for block in blocks:
                original       = (block.get("text") or "").strip()
                translated     = (block.get("translated_text") or original).strip()
                was_translated = block.get("was_translated", False)
                confidence     = block.get("confidence", 0.0)

                if not original:
                    continue

                if was_translated:
                    _add_translated_block(doc, original, translated, confidence)
                else:
                    _add_plain_block(doc, translated)

    # ── JSON appendix ─────────────────────────────────────────────────────────
    _add_metadata_appendix(doc, pages_data)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
